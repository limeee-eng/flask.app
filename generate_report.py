#!/usr/bin/env python3
"""生成业务逻辑及越权漏洞检测报告 (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

# ── 封面标题 ──────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('\n\n\n\n\n')
title.add_run('\n')

t = doc.add_paragraph()
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = t.add_run('业务逻辑漏洞及越权漏洞\n检测报告')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = sub.add_run('检测目标：用户信息管理平台\n检测日期：2026-07-09\n检测分支：vulnerable')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── 目录 ──────────────────────────────────────────────────
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

add_heading_styled('目录', 1)
toc_items = [
    '1. 报告概述',
    '2. 漏洞 1：IDOR 越权访问他人资料',
    '3. 漏洞 2：IDOR 越权为他人充值',
    '4. 漏洞 3：负金额充值（余额操纵）',
    '5. 漏洞 4：未登录访问绕过（仅 Session 校验）',
    '6. 漏洞 5：SQL 注入配合越权获取全量凭据',
    '7. 漏洞 6：余额无上限（金融逻辑缺陷）',
    '8. 漏洞 7：余额可为负值（会计逻辑缺陷）',
    '9. 漏洞 8：类型转换异常导致静默失败',
    '10. 漏洞矩阵与修复建议',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 辅助函数 ──────────────────────────────────────────────
def add_vuln_table(doc, rows_data):
    """添加漏洞详情表格"""
    table = doc.add_table(rows=len(rows_data) + 1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, text in enumerate(['项目', '内容']):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    # 数据行
    for idx, (key, val) in enumerate(rows_data):
        table.rows[idx + 1].cells[0].text = str(key)
        table.rows[idx + 1].cells[1].text = str(val)

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)

    return table


def add_code(doc, code):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


# ── 1. 概述 ──────────────────────────────────────────────
add_heading_styled('1. 报告概述', 1)
doc.add_paragraph(
    '本报告对用户信息管理平台的个人中心（/profile）和充值（/recharge）功能模块 '
    '进行了全面的业务逻辑漏洞和越权漏洞检测。检测发现共 8 类安全漏洞，'
    '其中严重级别 3 项，高危 3 项，中危 2 项。'
)

doc.add_paragraph()
add_heading_styled('检测范围', 2)
add_vuln_table(doc, [
    ('检测目标', '用户信息管理平台 — 个人中心 & 充值模块'),
    ('检测分支', 'vulnerable（漏洞版）'),
    ('检测日期', '2026-07-09'),
    ('检测工具', 'curl / 手动渗透测试'),
    ('登录用户', 'alice（普通用户, password=alice2025）'),
    ('测试端点', 'GET /profile, POST /recharge, GET /search'),
])

doc.add_page_break()

# ── 2. IDOR 越权访问他人资料 ─────────────────────────────
add_heading_styled('2. 漏洞 1：IDOR 越权访问他人资料', 1)

add_vuln_table(doc, [
    ('漏洞类型', 'IDOR（Insecure Direct Object Reference）'),
    ('风险等级', '🔴 严重'),
    ('利用难度', '极低'),
    ('漏洞位置', 'GET /profile?user_id={id}'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '/profile 路由从 URL 参数获取 user_id，直接查询并返回用户资料，'
    '未验证当前登录用户与所查询的 user_id 是否匹配。'
    '攻击者只需修改 URL 中的 user_id 参数值，即可查看任意用户的敏感信息。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞代码')
run.bold = True
add_code(doc, '@app.route("/profile")\n'
         'def profile():\n'
         '    if "username" not in session:\n'
         '        return redirect("/login")\n'
         '    try:\n'
         '        user_id = int(request.args.get("user_id", 0))\n'
         '    except ValueError:\n'
         '        user_id = 0\n'
         '    user_data = find_user_by_id(user_id)\n'
         '    # ❌ 未校验当前用户与 user_id 的匹配关系\n'
         '    ...')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True

doc.add_paragraph('以 alice 用户登录（普通用户），通过修改 URL 参数查看管理员 admin 的资料：')
add_code(doc, '# alice 查看 admin 的资料（user_id=1）\n'
         '$ curl -b cookies.txt "http://localhost:5000/profile?user_id=1"\n'
         '# ✅ 成功：获取到 admin 的邮箱、手机、余额\n'
         '# 泄露数据: admin@example.com, 13800138000, 余额 100499.0')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试结果')
run.bold = True
add_vuln_table(doc, [
    ('请求用户', 'alice（user_id=2, 普通用户）'),
    ('目标用户', 'admin（user_id=1, 管理员）'),
    ('修改参数', 'profile?user_id=1'),
    ('结果', '✅ 成功查看管理员全部资料'),
    ('泄露字段', '用户名、邮箱、手机、角色、余额'),
])

doc.add_page_break()

# ── 3. IDOR 越权充值 ─────────────────────────────────────
add_heading_styled('3. 漏洞 2：IDOR 越权为他人充值', 1)

add_vuln_table(doc, [
    ('漏洞类型', 'IDOR 越权操作'),
    ('风险等级', '🔴 严重'),
    ('利用难度', '极低'),
    ('漏洞位置', 'POST /recharge'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '/recharge 路由从表单参数接收 user_id 和 amount，直接修改对应用户的余额。'
    '不校验当前登录用户与目标 user_id 的归属关系，'
    '导致任意登录用户可为任意账户充值（或通过负金额扣款）。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '# alice 给 admin（user_id=1）充值 500 元\n'
         '$ curl -b cookies.txt -d "user_id=1&amount=500" \\\n'
         '  http://localhost:5000/recharge\n'
         '# ✅ 充值成功，admin 余额从 99999 → 100499')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('风险说明')
run.bold = True
doc.add_paragraph(
    '- 攻击者可控制任意账户余额\n'
    '- 与负金额结合，可清空任意账户\n'
    '- 可用于洗钱、资金转移\n'
    '- 无法追溯操作者'
)

doc.add_page_break()

# ── 4. 负金额充值 ────────────────────────────────────────
add_heading_styled('4. 漏洞 3：负金额充值（余额操纵）', 1)

add_vuln_table(doc, [
    ('漏洞类型', '业务逻辑缺陷 — 输入校验缺失'),
    ('风险等级', '🔴 严重'),
    ('利用难度', '极低'),
    ('漏洞位置', 'POST /recharge, amount 参数'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    'amount 参数未做正负校验，攻击者可传入负值实现从任意账户扣款。'
    '结合 IDOR 漏洞，可清空任意用户余额。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞代码')
run.bold = True
add_code(doc, 'amount = float(request.form.get("amount", 0))\n'
         '# ❌ 未校验 amount > 0\n'
         'u["balance"] = u["balance"] + amount  # amount 可为负数')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '# alice 给自己充 -1000 元（实际是扣款）\n'
         '$ curl -b cookies.txt -d "user_id=2&amount=-1000" \\\n'
         '  http://localhost:5000/recharge\n'
         '# ✅ 成功，alice 余额从 100 → -900\n\n'
         '# 结合 IDOR 可清空 admin 账户\n'
         '$ curl -b cookies.txt -d "user_id=1&amount=-100000" \\\n'
         '  http://localhost:5000/recharge')

doc.add_page_break()

# ── 5. 未登录访问 ────────────────────────────────────────
add_heading_styled('5. 漏洞 4：未登录访问绕过（仅 Session 校验）', 1)

add_vuln_table(doc, [
    ('漏洞类型', '认证绕过'),
    ('风险等级', '🟡 中危'),
    ('利用难度', '低'),
    ('漏洞位置', 'GET /profile, POST /recharge'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '/profile 和 /recharge 仅通过 session 中的 username 判断登录状态，'
    '未对用户权限做分级校验。一旦 session 被劫持或伪造，攻击者可完全控制所有功能。'
    '此外，未登录访问时返回 302 跳转，但未使用 HTTP 401/403 状态码。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '$ curl -s -o /dev/null -w "HTTP %{http_code}" \\\n'
         '  "http://localhost:5000/profile?user_id=1"\n'
         '# HTTP 302（重定向到 /login，非 401/403）')

doc.add_page_break()

# ── 6. SQL注入 + 越权 ────────────────────────────────────
add_heading_styled('6. 漏洞 5：SQL 注入配合越权获取全量凭据', 1)

add_vuln_table(doc, [
    ('漏洞类型', 'SQL 注入 + 数据泄露'),
    ('风险等级', '🔴 严重'),
    ('利用难度', '低'),
    ('漏洞位置', 'GET /search?keyword='),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '搜索接口使用 f-string 拼接 SQL，且 session 中用户登录后即可访问。'
    '攻击者通过 UNION 注入可脱取数据库中所有用户的用户名、密码明文、邮箱、手机号。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, "# 脱取所有用户凭据（含密码）\n"
         "$ curl -b cookies.txt -G 'http://localhost:5000/search' \\\n"
         "  --data-urlencode \"keyword=' UNION SELECT id,username,password,email,phone FROM users --\"\n"
         "# 结果：admin/admin123, alice/alice2025 等全部泄露")

add_vuln_table(doc, [
    ('泄露的用户数', '5+（含之前注入创建的测试用户）'),
    ('泄露的密码', 'admin123, alice2025 等明文密码'),
    ('泄露的敏感字段', '用户名、密码、邮箱、手机号'),
    ('影响', '全量用户凭据泄露，可进一步社工或撞库攻击'),
])

doc.add_page_break()

# ── 7. 余额无上限 ────────────────────────────────────────
add_heading_styled('7. 漏洞 6：余额无上限（金融逻辑缺陷）', 1)

add_vuln_table(doc, [
    ('漏洞类型', '业务逻辑缺陷'),
    ('风险等级', '🟡 中危'),
    ('利用难度', '低'),
    ('漏洞位置', 'POST /recharge'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '充值接口未对金额上限做任何限制。攻击者可一次充值极大金额（如 999999999），'
    '导致余额出现天文数字，破坏系统数据一致性和财务报表。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '# alice 充 999999999 元\n'
         '$ curl -b cookies.txt -d "user_id=2&amount=999999999" \\\n'
         '  http://localhost:5000/recharge\n'
         '# ✅ 成功，alice 余额变为 999999099.01\n'
         '# ⚠️ 余额无上限，可导致数据溢出或精度丢失')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('浮点数精度问题')
run.bold = True
doc.add_paragraph(
    'amount 使用 float 类型存储和计算，存在浮点数精度问题。'
    '累积多次 0.01 元的充值可能导致余额与实际不符，'
    '在金融场景中可能被利用进行"精度套利"。'
)

doc.add_page_break()

# ── 8. 余额可为负值 ──────────────────────────────────────
add_heading_styled('8. 漏洞 7：余额可为负值（会计逻辑缺陷）', 1)

add_vuln_table(doc, [
    ('漏洞类型', '业务逻辑缺陷'),
    ('风险等级', '🟡 中危'),
    ('利用难度', '低'),
    ('漏洞位置', 'POST /recharge + GET /profile'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    '未校验余额扣减后是否为负值，导致用户余额可为负数（-900）。'
    '攻击者可利用此漏洞无限透支，造成系统金融数据混乱。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '# alice 余额原为 100，充值 -1000\n'
         '$ curl -b cookies.txt -d "user_id=2&amount=-1000" \\\n'
         '  http://localhost:5000/recharge\n'
         '# ✅ 余额变为 -900（负值）')

doc.add_page_break()

# ── 9. 类型转换异常 ──────────────────────────────────────
add_heading_styled('9. 漏洞 8：类型转换异常导致静默失败', 1)

add_vuln_table(doc, [
    ('漏洞类型', '异常处理缺陷'),
    ('风险等级', '🟢 低危'),
    ('利用难度', '低'),
    ('漏洞位置', 'GET /profile, POST /recharge'),
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('漏洞描述')
run.bold = True
doc.add_paragraph(
    'user_id 和 amount 参数使用 int() / float() 强制类型转换，'
    '异常时静默降级为默认值（user_id=0, amount=0），'
    '而非返回明确的错误提示。这种静默失败可能导致用户混淆，'
    '且与搜索功能中的 SQL 注入配合时，攻击者难以被及时发现。'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('测试记录')
run.bold = True
add_code(doc, '# 字符串 user_id → 降级为 0 → 显示"用户不存在"\n'
         '$ curl -b cookies.txt "http://localhost:5000/profile?user_id=abc"\n'
         '# 结果："用户不存在"\n\n'
         '# 非数字 amount → 降级为 0 → 重定向到 profile\n'
         '$ curl -b cookies.txt -d "user_id=2&amount=abc" \\\n'
         '  http://localhost:5000/recharge\n'
         '# 结果：HTTP 302 重定向，无错误提示')

doc.add_page_break()

# ── 10. 漏洞矩阵与修复建议 ──────────────────────────────
add_heading_styled('10. 漏洞矩阵与修复建议', 1)

add_heading_styled('漏洞矩阵', 2)

# 创建漏洞总结表
table = doc.add_table(rows=9, cols=5, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['#', '漏洞名称', '严重程度', '利用难度', '影响']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

vuln_data = [
    ('1', 'IDOR 越权查看他人资料', '🔴 严重', '极低', '全量用户信息泄露'),
    ('2', 'IDOR 越权为他人充值', '🔴 严重', '极低', '资金数据篡改'),
    ('3', '负金额充值（余额操纵）', '🔴 严重', '极低', '任意扣款/透支'),
    ('4', '未登录访问仅 Session 校验', '🟡 中危', '低', '认证机制薄弱'),
    ('5', 'SQL 注入 + 越权数据脱取', '🔴 严重', '低', '全库凭据泄露'),
    ('6', '余额无上限', '🟡 中危', '低', '数据一致性破坏'),
    ('7', '余额可为负值', '🟡 中危', '低', '会计数据异常'),
    ('8', '类型转换异常静默失败', '🟢 低危', '低', '用户体验缺陷'),
]

for idx, (num, name, severity, difficulty, impact) in enumerate(vuln_data):
    row = table.rows[idx + 1]
    row.cells[0].text = num
    row.cells[1].text = name
    row.cells[2].text = severity
    row.cells[3].text = difficulty
    row.cells[4].text = impact

# 设置列宽
for row in table.rows:
    row.cells[0].width = Cm(1)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(2)
    row.cells[4].width = Cm(4)

doc.add_paragraph()

add_heading_styled('修复建议', 2)

add_heading_styled('🔴 紧急修复', 3)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('1. 修复 IDOR 越权 — 校验资源归属')
run.bold = True
add_code(doc,
         '# profile 中校验当前用户与目标用户匹配\n'
         'def profile():\n'
         '    if "username" not in session:\n'
         '        return redirect("/login")\n'
         '    current_user = session["username"]\n'
         '    try:\n'
         '        user_id = int(request.args.get("user_id", 0))\n'
         '    except ValueError:\n'
         '        abort(400)\n'
         '    # ✅ 只允许查看自己的资料\n'
         '    if USERS[current_user]["id"] != user_id:\n'
         '        abort(403)')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('2. 修复负金额漏洞 — 校验 amount > 0')
run.bold = True
add_code(doc,
         '# recharge 中校验金额必须为正\n'
         'def recharge():\n'
         '    amount = float(request.form.get("amount", 0))\n'
         '    # ✅ 金额必须为正\n'
         '    if amount <= 0:\n'
         '        return render_template("error.html",\n'
         '                                error="金额必须大于 0")\n'
         '    # ✅ 金额上限\n'
         '    if amount > 100000:\n'
         '        return render_template("error.html",\n'
         '                                error="单次充值不能超过 100,000 元")')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('3. 修复余额负值 — 校验扣减后余额')
run.bold = True
add_code(doc,
         '# 仅允许管理员扣款，且余额不能为负\n'
         'new_balance = u["balance"] + amount\n'
         'if new_balance < 0:\n'
         '    return render_template("error.html",\n'
         '                            error="余额不足")\n'
         'u["balance"] = new_balance')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('4. 修复 SQL 注入 — 使用参数化查询')
run.bold = True
add_code(doc,
         '# search 中使用参数化查询替代 f-string\n'
         'cursor.execute(\n'
         '    "SELECT id, username, email, phone "\n'
         '    "FROM users WHERE username LIKE ? OR email LIKE ?",\n'
         '    (f"%{keyword}%", f"%{keyword}%")\n'
         ')')

doc.add_paragraph()

add_heading_styled('🟡 建议修复', 3)

doc.add_paragraph('5. 添加操作审计日志 — 记录谁在何时对哪个账户做了什么操作', style='List Bullet')
doc.add_paragraph('6. 添加余额变动上限 — 单日充值/扣款上限', style='List Bullet')
doc.add_paragraph('7. 使用 Decimal 替代 Float — 避免浮点数精度问题', style='List Bullet')
doc.add_paragraph('8. 统一错误处理 — 类型转换失败返回 400 Bad Request', style='List Bullet')
doc.add_paragraph('9. Session 加固 — 增加 User-Agent 绑定、IP 绑定', style='List Bullet')
doc.add_paragraph('10. 速率限制 — 充值接口限制频率，防止批量攻击', style='List Bullet')

doc.add_paragraph()

# ── 免责声明 ──────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('免责声明：本报告仅用于安全教育和漏洞验证。请在获得授权的前提下进行安全测试。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── 保存 ──────────────────────────────────────────────────
output_path = '/home/user/user-management/业务逻辑漏洞及越权漏洞检测报告.docx'
doc.save(output_path)
print(f'报告已保存到: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
